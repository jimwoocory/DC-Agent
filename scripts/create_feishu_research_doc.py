from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("巅池文化全流程与飞书多维表格落地调研报告.docx")
FONT = "Microsoft YaHei"

BLACK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D9E2EC"


def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size: float = 10.5, color: RGBColor = BLACK):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="B8C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    size = {1: 16, 2: 13, 3: 12}.get(level, 11)
    set_run_font(run, size=size, bold=True, color=BLUE if level < 3 else BLACK)
    return p


def add_body(doc, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True, color=BLACK)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, size=10.5, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=BLACK)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)


def add_callout(doc, label: str, text: str, fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="D3DFEA", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, size=10.5, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr_cells[idx], header_fill)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=BLACK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx in (0, 1):
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if len(str(value)) <= 8
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)


def add_masthead(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("巅池文化全流程与飞书多维表格落地调研报告")
    set_run_font(run, size=22, bold=True, color=BLACK)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run("阶段性底稿｜围绕执行部与品宣部流程图判断飞书能不能落地")
    set_run_font(r, size=12, color=MUTED)

    meta = [
        ("汇报对象", "杨总"),
        ("文档定位", "把流程图转成可跟进、可汇总、可验收的管理工具"),
        ("当前状态", "已查后台当前版本；待飞书官方确认升级费用和落地方式"),
        ("日期", "2026 年 5 月 21 日"),
    ]
    add_simple_table(
        doc, ["项目", "内容"], meta, [1800, 7560], header_fill="FFFFFF", font_size=9.5
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    setup_styles(doc)

    header_p = section.header.paragraphs[0]
    header_p.text = "飞书能力与升级方案调研"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header_p, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "内部讨论稿｜报价和能力以飞书官方书面回复为准"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer_p, size=9, color=MUTED)

    add_masthead(doc)

    add_callout(
        doc,
        "一句话结论",
        "杨总决定看多维表格，本质是想把执行部、品宣部两张流程图，从 XMind 变成公司每天能用的管理底账。重点不是软件功能多不多，而是项目、活动、内容、直播、舆情、礼品、库存、验收这些事，能不能有人负责、有节点、有提醒、有总表。",
    )

    add_heading(doc, "一、录音里杨总真正关注的点", 1)
    add_body(
        doc,
        "这次汇报不能写成普通软件介绍。杨总真正问的是：多维表格现在看到的只是一些统计图，如果看不到具体项目、具体责任人、具体节点和具体材料，那对业务管理意义不够。",
    )
    add_simple_table(
        doc,
        ["杨总表达的意思", "要落实成什么问题"],
        [
            (
                "先搞清楚用的是飞书哪一类、哪个应用",
                "确认当前版本、当前开通能力，以及多维表格到底处在哪个套餐权限里",
            ),
            (
                "多维表格特点是什么，功能是什么，能不能满足需求",
                "不是看功能列表，而是看能不能承接执行部、品宣部流程",
            ),
            (
                "现在能看到总表/数据台，但没有具体内容",
                "确认为什么只能看到柱状图或汇总数据，看不到下钻明细",
            ),
            (
                "权限不够，那给这个数据台意义是什么",
                "确认升级后是否能看到项目明细、责任人、流程节点、材料状态",
            ),
            (
                "钱能解决的问题不是问题",
                "价格不是第一位，关键是多花钱后能不能把管理问题解决",
            ),
            (
                "明确，不要模棱两可",
                "必须拿到飞书官方明确答复：能不能做、怎么做、多少钱、多久",
            ),
        ],
        [3600, 5760],
        font_size=8.6,
    )

    add_heading(doc, "二、目前后台已经确认的版本情况", 1)
    add_simple_table(
        doc,
        ["项目", "后台显示结果", "对这次判断的影响"],
        [
            (
                "当前产品",
                "飞书商业标准版",
                "不是免费版，但也不是更高等级版本；部分管理、数据和自动化能力需要确认是否够用",
            ),
            (
                "服务时间",
                "2026-01-09 ~ 2027-01-08",
                "如果中途升级，需要问清楚补差价和续费口径",
            ),
            (
                "席位情况",
                "已购 34 个，已分配 26 个",
                "当前人数够用；重点不是席位，而是能力是否够",
            ),
            (
                "组织人数",
                "26 人",
                "流程落地涉及市场、策划、执行、设计、影视、运营、媒介、社群、财务等多角色",
            ),
            (
                "多维表格自动化",
                "43 次 / 200 次",
                "如果流程提醒、到期提醒、状态变更都用自动化，200 次/月可能很快不够",
            ),
            (
                "多维表格行数",
                "行数用量页显示多张表为 20,000 行上限；套餐对比显示商业标准版基础权益为 2,000 行",
                "这里存在限免/扩容口径，需要飞书确认长期正式上限",
            ),
            (
                "数据表扩容",
                "后台显示每张表可扩容上限可在 0 至 198 万行之间调整",
                "行数容量可能不是第一卡点，但扩容是否收费、是否长期有效要确认",
            ),
            (
                "集成平台运行",
                "0 次 / 500 次",
                "能做部分自动流转，但能否连接外部知识库/NAS需确认",
            ),
            (
                "在线数据源数据条数",
                "当前套餐显示为 0 条",
                "如果要接外部数据源或做公司知识库同步，当前版本可能卡住",
            ),
            (
                "知识库使用数据",
                "后台显示为增值能力/升级咨询页面",
                "Wiki 相关数据看板和管理能力当前受限，需要官方确认开通方式",
            ),
        ],
        [1700, 3100, 4560],
        font_size=8.2,
    )

    add_heading(doc, "三、多维表格当前最关键的限制", 1)
    add_callout(
        doc,
        "业务判断",
        "现在不是完全不能用多维表格，而是当前后台更多只能看到用量、趋势、柱状图、行数、流程次数这类管理数据；杨总真正要的是能下钻到业务明细：哪个项目、哪一步、谁负责、什么材料没交、为什么延期。",
        fill=LIGHT_GRAY,
    )
    add_simple_table(
        doc,
        ["后台/能力", "现在能看到什么", "杨总真正需要什么"],
        [
            (
                "云文档功能使用情况",
                "活跃人数、创建文件人数、创建文件数量等趋势图",
                "具体哪些业务资料、哪些项目文件、谁负责更新、是否可追溯",
            ),
            (
                "应用使用数据",
                "页面显示为增值能力，可看应用安装、可用人数、使用人数和分部门详情",
                "如果要看谁在用、哪个部门没用起来，需要确认是否必须升级",
            ),
            (
                "多维表格流程用量",
                "能看到流程名称、关联多维表格、最后编辑人、运行次数",
                "流程是否对应真实业务节点，是否能提醒到责任人，是否能追踪卡点",
            ),
            (
                "多维表格行数用量",
                "能看到数据表、关联多维表格、所有者、已使用行数/上限",
                "行数够不够不是唯一问题，更重要是能不能按项目、部门、阶段看明细",
            ),
            (
                "多维表格仪表盘/总览",
                "可以做图表和汇总，但当前反馈是只能看到简单柱状数据",
                "图表必须能回到明细表，看到项目、负责人、节点、材料和问题原因",
            ),
            (
                "权限能力",
                "后台套餐对比显示高级视图、列权限控制等多维能力与套餐相关",
                "市场、策划、执行、设计、影视、运营、财务要各看各的，又能给老板汇总",
            ),
        ],
        [2100, 3300, 3960],
        font_size=8.1,
    )

    add_heading(doc, "四、杨总真正要的结果", 1)
    for item in [
        "把执行部和品宣部流程图落成实际工具，而不是只停留在 XMind 上。",
        "每个项目走到哪一步、谁负责、卡在哪里，杨总能一眼看到。",
        "方案、设计、场地、物料、视频、礼品、直播、内容、舆情、社群、库存、验收材料都有对应记录，不靠人反复追问。",
        "如果多维表格能做，明确要怎么搭、多少钱、多久能用起来。",
        "如果多维表格做不到，明确用飞书其他工具补，还是准备看外部平台。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "五、两张流程图里真正要管住的业务", 1)
    add_simple_table(
        doc,
        ["部门/流程线", "要管住什么", "多维表格里应该看到什么"],
        [
            (
                "执行部-活动执行",
                "甲方需求、方案审核、设计、场地、项目启动会、布置落地、点检、活动结束、复盘验收",
                "项目总表、阶段状态、负责人、截止时间、风险提醒、验收材料清单",
            ),
            (
                "执行部-礼品生产",
                "设计文件、供应商打样、生产合同、大货进度、抽检、延期处理、交付验收",
                "打样进度表、生产进度表、费用与时间同步、异常处理记录",
            ),
            (
                "执行部-出入库",
                "礼品库存、入库、出库、财务核对、库存责任",
                "库存表、入库记录、出库记录、每周库存同步、财务核对记录",
            ),
            (
                "品宣部-直播运营",
                "新旧账号、运营目标、上架直播、数据跟踪、周/月复盘、爆品和选品调整",
                "账号表、直播计划、数据复盘表、KPI达成、爆品记录",
            ),
            (
                "品宣部-账号内容",
                "账号定位、内容发布、数据跟踪、复盘改进、舆情投诉、好事分享、直连获客、社群转化",
                "内容日历、账号复盘、线索转化、闭环记录",
            ),
            (
                "品宣部-舆情社群",
                "用户言论、素材收集、风险预警、负面稀释、信息扩散、建群维护、用户邀约",
                "舆情台账、素材库、社群管理表、预警报告、复盘交付",
            ),
            (
                "品宣部-媒介/KOC/用户故事",
                "媒介库、KOC库、用户故事匹配、拍摄、成片交付、验收材料",
                "资源库、需求下发表、日报复盘、用户故事进度",
            ),
        ],
        [1900, 4300, 3160],
        font_size=8.4,
    )

    add_heading(doc, "六、建议先搭的多维表格底账", 1)
    add_body(
        doc,
        "两张流程图都很完整，但不能一口气做成很复杂的系统。建议先按“项目总表统管，部门业务表配合”的方式搭底账，先把关键节点跑起来。",
    )
    add_simple_table(
        doc,
        ["表/视图", "用途", "杨总能看到的结果"],
        [
            (
                "项目总表",
                "记录公司每个甲方需求、项目或运营任务的全过程",
                "所有项目进度、负责人、卡点、是否延期",
            ),
            (
                "方案与审核表",
                "记录需求分析、方案输出、内部审核、甲方反馈、通过/不通过",
                "哪些方案没过、为什么没过、改到哪一步",
            ),
            (
                "设计制作表",
                "记录设计需求、初稿、策划审核、甲方修改、制作文件入库、物料工厂",
                "设计是否拖进度，制作文件是否齐",
            ),
            (
                "场地与落地表",
                "记录勘场、场地确认、合同、服务、落地沟通会、现场点位",
                "场地是否确定，落地前事项是否对齐",
            ),
            (
                "物料安装表",
                "记录报价、合规、中标、制作、安装、验收、财务验收",
                "物料安装是否按时、安全、合规",
            ),
            (
                "视频执行表",
                "记录分镜、拍摄通告、机位、现场拍摄、后期剪辑、成片交付",
                "视频交付是否跟上活动节奏",
            ),
            (
                "礼品生产表",
                "记录打样、供应商、合同、大货、抽检、延期、交付验收",
                "礼品是否按时合格交付",
            ),
            (
                "库存出入库表",
                "记录库存、入库、出库、财务核对和责任人",
                "库存是否准确，责任是否清楚",
            ),
            (
                "账号运营表",
                "记录直播账号、内容账号的新旧状态、目标、发布、数据、复盘",
                "KPI是否达成，爆品和问题是否沉淀",
            ),
            (
                "舆情社群表",
                "记录用户言论、风险预警、负面稀释、信息扩散、社群维护",
                "风险是否及时发现，处理是否闭环",
            ),
            (
                "资源库",
                "记录媒介、KOC、用户故事、素材、竞品案例等可复用资源",
                "资源是否能被市场和策划快速调用",
            ),
            (
                "验收复盘表",
                "记录验收材料、成本票据、问题复盘、下次改进",
                "项目结束后有没有闭环",
            ),
        ],
        [1850, 4300, 3210],
        font_size=8.2,
    )

    add_heading(doc, "七、杨总需要的几个关键页面", 1)
    for item in [
        "老板总览：所有项目现在处于什么阶段，哪些正常，哪些延期，哪些需要协调。",
        "负责人工作台：每个人今天要处理什么，哪些已经超期，哪些等别人确认。",
        "项目详情页：单个项目从需求到验收的完整记录，避免信息散在聊天里。",
        "风险提醒页：方案未通过、设计反复改、场地未定、物料延期、礼品不合格、舆情风险等问题集中显示。",
        "运营复盘页：直播、内容、账号、社群每周/月复盘后，能看到达成KPI、爆品、延展机会和改进动作。",
        "验收材料页：活动结束后，需要交市场部和财务的材料是否齐全。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "八、知识库与资料沉淀风险", 1)
    add_callout(
        doc,
        "关键风险",
        "如果流程资料、项目文件、复盘结果都放进飞书，但公司自己的 DC-Agent、NAS、AstrBot 知识库读不出来，资料还是被锁在飞书里。这样短期能协作，长期沉淀和 AI 查询会受限制。",
        fill=LIGHT_GRAY,
    )
    add_simple_table(
        doc,
        ["要解决的事", "现在的风险", "要向飞书确认的口径"],
        [
            (
                "云文档读取",
                "公司云文档如果没有接口权限或文档权限，DC-Agent 不能稳定读取",
                "当前商业标准版是否支持按公司授权读取云文档内容",
            ),
            (
                "Wiki 读取",
                "后台知识库相关数据能力显示受限，实际内容读取也可能受权限影响",
                "Wiki 是否需要单独开通；开通后能否同步到自建知识库",
            ),
            (
                "多维表格同步",
                "流程表建起来后，如果跨表汇总、附件、历史记录不能稳定导出，会影响复盘",
                "多维表格记录、附件、视图、权限能否通过官方方式读取",
            ),
            (
                "NAS 归档",
                "资料不能定期同步出来，就无法形成公司自己的资料资产",
                "是否支持定期导出或接口同步到 NAS",
            ),
            (
                "AstrBot 知识库",
                "资料进不了知识库，AI 只能回答零散内容，不能沉淀公司经验",
                "飞书官方推荐用哪种方式接入外部知识库",
            ),
        ],
        [1800, 3800, 3760],
        font_size=8.4,
    )

    add_heading(doc, "九、向飞书官方必须问清楚的问题", 1)
    add_body(
        doc,
        "以下问题建议直接发给飞书客户经理，要求对方围绕两张流程图演示，不要只讲产品介绍。",
    )
    official_questions = [
        "按《执行部全流程》和《品宣部全流程》，多维表格能不能搭出项目总表和各业务明细表？",
        "为什么当前只能看到简单柱状图/汇总图，看不到业务明细？这是版本限制、权限设置问题，还是当前表没有搭对？",
        "升级或开通对应能力后，仪表盘能不能下钻到具体项目、负责人、节点、材料和问题原因？",
        "一个项目从甲方需求到验收复盘，能不能自动带出对应的阶段、负责人、截止时间和待办？",
        "老板总览能不能看到所有项目进度、运营KPI、延期、风险、成本票据和验收材料状态？",
        "员工填表后，相关负责人能不能收到提醒？超期未处理能不能提醒？",
        "项目中的方案、设计稿、场地图、合同、验收材料，能不能统一挂在对应项目下面？",
        "市场、策划、执行、设计、影视、运营、媒介、社群、财务等角色，能不能做到各看各的、各填各的、负责人可汇总？",
        "直播数据、内容数据、舆情记录、用户故事、媒介/KOC日报，能不能按周/月复盘并沉淀？",
        "公司云文档、Wiki、多维表格内容，能不能通过官方方式同步到公司 NAS 和 AstrBot 知识库？",
        "当前商业标准版能不能做资料同步？如果不能，升级到哪个版本能做？",
        "要实现资料同步，需要开哪些权限、配哪些应用、是否需要飞书协助配置？",
        "如果不能开放接口，是否支持定期完整导出，保证公司资料不被锁在飞书里？",
        "现有版本能做到哪些？哪些必须升级？升级后能多做到哪些？",
        "如果要飞书协助搭建这套流程，费用多少、周期多久、交付到什么程度？",
        "如果多维表格不够，飞书建议用多维表格加妙搭、项目、审批，还是用其他组合？",
        "能不能拿我们这张流程图做一次演示，现场跑一个项目样例，看是否真的适合？",
    ]
    for q in official_questions:
        add_number(doc, q)

    add_heading(doc, "十、版本与方案对比表", 1)
    add_body(
        doc, "这张表先作为询价和判断模板，价格和结论不要猜，等飞书官方回复后补齐。"
    )
    add_simple_table(
        doc,
        ["方案", "能解决什么", "还要确认什么", "费用", "初步判断"],
        [
            (
                "当前商业标准版",
                "基础协同、基础表格、部分自动化",
                "流程提醒、跨表汇总、Wiki、外部同步是否被限制",
                "已购",
                "可以先试搭，但可能做不完整",
            ),
            (
                "升级飞书版本",
                "项目总表、阶段跟进、负责人待办、资料汇总、更多自动化可能更完整",
                "升级后能否支撑老板总览、多角色协同、NAS/AstrBot同步",
                "待确认",
                "优先确认，可能是最快方案",
            ),
            (
                "多维表格+飞书其他工具",
                "复杂审批、项目协同、流程提醒、知识库沉淀可能更完整",
                "组合后员工使用是否顺手，资料能否同步出来",
                "待确认",
                "作为飞书内部加强版",
            ),
            (
                "外部平台",
                "可能更适合做流程系统和知识库连接",
                "迁移成本、员工使用成本、和飞书配合程度",
                "待确认",
                "飞书不满足时再深入比较",
            ),
        ],
        [1500, 2450, 2500, 1200, 1710],
        font_size=8.7,
    )

    add_heading(doc, "十一、外部竞品横向比较方向", 1)
    add_body(
        doc,
        "外部平台不建议一开始就深入采购，只先做备选池。等飞书官方明确答复后，再判断有没有必要进一步约演示。",
    )
    add_simple_table(
        doc,
        ["平台方向", "适合看什么", "重点风险", "是否进入下一轮"],
        [
            (
                "简道云/明道云/轻流",
                "按公司流程搭管理工具，适合非技术人员维护",
                "费用、后期维护、和飞书配合程度",
                "待定",
            ),
            (
                "伙伴云/维格表等表格型工具",
                "多表协作、数据汇总、看板展示",
                "能否承接完整流程，不只是做表",
                "待定",
            ),
            (
                "定制系统",
                "按公司习惯完全定做",
                "周期长、费用高、后期要人维护",
                "最后选项",
            ),
        ],
        [2200, 2900, 2800, 1460],
        font_size=9,
    )

    add_heading(doc, "十二、建议的判断标准", 1)
    add_callout(
        doc,
        "拍板标准",
        "只看一个问题：这套方案能不能把两张流程图变成每天可执行、可沉淀、可复盘的管理工具。如果能让项目少漏事、运营有复盘、风险有提醒、资料能进公司知识库，才谈价格值不值；如果不能，再便宜也不是合适方案。",
        fill=LIGHT_GRAY,
    )
    for item in [
        "能不能承接两张流程图里的真实节点，而不是让业务迁就工具。",
        "老板要看的项目进度、运营结果、延期、风险、验收结果能不能汇总出来，并且能点回明细。",
        "飞书里的资料能不能同步到 NAS 和 AstrBot，变成公司自己的知识库。",
        "市场、策划、执行、设计、影视、运营、媒介、社群、财务日常使用是否顺手。",
        "能不能先拿一个真实项目试跑，短期内看到效果。",
        "费用是否清楚：一次性费用、年费、服务费、后期维护费都要列出来。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十三、下一步行动表", 1)
    add_simple_table(
        doc,
        ["动作", "负责人", "产出", "时间"],
        [
            (
                "把两张流程图拆成表格清单",
                "我方",
                "项目总表和各业务表字段初稿",
                "1 天内",
            ),
            (
                "确认当前飞书版本和已购模块",
                "我方",
                "当前版本截图/合同/后台信息",
                "1 天内",
            ),
            (
                "向飞书官方发两张流程图和问题清单",
                "我方",
                "官方书面回复或会议纪要",
                "1-2 天",
            ),
            (
                "确认总览能否下钻明细",
                "飞书/我方",
                "柱状图、仪表盘、项目明细之间能否打通",
                "本周内",
            ),
            ("约飞书用真实项目演示", "飞书/我方", "是否能落地的明确判断", "本周内"),
            (
                "确认资料同步能力",
                "飞书/我方",
                "云文档、Wiki、多维表格同步到 NAS/AstrBot 的可行方案",
                "本周内",
            ),
            ("补齐费用表", "我方", "升级/定制/续费费用", "拿到报价当天"),
            ("形成推荐结论", "我方", "给杨总的汇报版 Word/PPT", "下周前"),
        ],
        [2300, 1200, 3300, 2560],
        font_size=9,
    )

    add_heading(doc, "十四、给杨总汇报时的建议说法", 1)
    add_body(
        doc,
        "建议开场：",
        bold_prefix="建议开场：",
    )
    add_body(
        doc,
        "杨总，这次我不把它当软件功能来查，我就按执行部和品宣部两张流程图来判断：多维表格能不能把项目进度、运营目标、负责人、节点提醒、资料验收和复盘结果管起来。现在最大的问题是，后台看到的只是柱状图和汇总数据还不够，关键要能点回明细，看到哪个项目、谁负责、卡在哪一步、材料交没交。我会让飞书官方按我们的流程跑一个样例，明确能不能做、要多少钱、多久能用；如果飞书做不到，我再把外部平台备选拉出来比较。",
    )

    add_heading(doc, "十五、后续 PPT 页结构", 1)
    ppt_pages = [
        ("第 1 页", "结论先行：两张流程图能不能落成管理工具"),
        ("第 2 页", "执行部和品宣部流程主线"),
        ("第 3 页", "当前多维表格限制：有汇总，缺明细"),
        ("第 4 页", "建议先搭的公司管理底账"),
        ("第 5 页", "杨总需要看的关键页面"),
        ("第 6 页", "飞书官方需要确认的问题"),
        ("第 7 页", "当前版本/升级/飞书组合方案对比"),
        ("第 8 页", "知识库与资料沉淀风险"),
        ("第 9 页", "费用、周期、落地风险"),
        ("第 10 页", "推荐方案与下一步动作"),
    ]
    add_simple_table(doc, ["页码", "内容"], ppt_pages, [1500, 7860], font_size=9.5)

    add_heading(doc, "十六、已核实依据", 1)
    add_body(
        doc, "以下信息用于内部判断，最终报价和开通范围仍以飞书客户经理书面回复为准。"
    )
    add_simple_table(
        doc,
        ["依据", "已看到/已确认的信息"],
        [
            (
                "飞书管理后台-我的产品",
                "当前产品为飞书商业标准版；服务期 2026-01-09 ~ 2027-01-08；已购 34 席，已分配 26 席",
            ),
            (
                "飞书管理后台-权益数据",
                "多维表格自动化 43/200 次；总存储 504.21GB/1.46TB；API 调用次数不限",
            ),
            (
                "飞书管理后台-多维表格流程用量",
                "能看到流程名称、关联多维表格、最后编辑人、运行次数；当前有 4 条流程记录",
            ),
            (
                "飞书管理后台-多维表格行数用量",
                "能看到数据表名称、关联多维表格、所有者、已使用行数/行数上限；部分表显示 20,000 行上限",
            ),
            (
                "飞书管理后台-应用使用数据",
                "该页面为增值能力，介绍可按单个应用查看分部门详情并导出；当前套餐页显示需要升级咨询",
            ),
            (
                "飞书管理后台-云文档功能使用情况",
                "当前可看到云文档活跃人数、创建文件人数、创建文件数量等趋势图",
            ),
            (
                "飞书管理后台-知识库使用数据",
                "该页面显示升级咨询/套餐对比，说明知识库使用数据看板属于受限能力",
            ),
            (
                "飞书开放平台文档",
                "云文档、多维表格、附件、权限等接口都需要对应应用权限和资料本身权限，不是管理员后台登录就能直接读全部资料",
            ),
        ],
        [2600, 6760],
        font_size=8.7,
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
